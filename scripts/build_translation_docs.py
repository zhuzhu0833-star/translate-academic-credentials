#!/usr/bin/env python3
"""Build English-only credential translation as .docx and .pdf from JSON."""

from __future__ import annotations

import json
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Any

ALIGNMENTS = {
    "left": "LEFT",
    "center": "CENTER",
    "right": "RIGHT",
    "justify": "JUSTIFY",
}


def load_input(path: Path) -> dict[str, Any]:
    with path.open(encoding="utf-8") as f:
        return json.load(f)


def get_docx_alignment(block: dict[str, Any], default: str = "left") -> Any:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    name = block.get("align", default).lower()
    return getattr(WD_ALIGN_PARAGRAPH, ALIGNMENTS.get(name, "LEFT"))


def apply_paragraph_style(paragraph: Any, block: dict[str, Any]) -> None:
    from docx.shared import Pt

    fmt = paragraph.paragraph_format
    if "space_before" in block:
        fmt.space_before = Pt(block["space_before"])
    if "space_after" in block:
        fmt.space_after = Pt(block["space_after"])
    paragraph.alignment = get_docx_alignment(block)

    if not paragraph.runs:
        paragraph.add_run("")
    run = paragraph.runs[0]
    if block.get("bold"):
        run.bold = True
    if "font_size" in block:
        run.font.size = Pt(block["font_size"])


def add_spacer(doc: Any, lines: int) -> None:
    for _ in range(max(1, lines)):
        doc.add_paragraph("")


def add_paragraph_block(doc: Any, block: dict[str, Any]) -> None:
    paragraph = doc.add_paragraph(block.get("text", ""))
    apply_paragraph_style(paragraph, block)


def add_field_row_block(doc: Any, block: dict[str, Any]) -> None:
    separator = block.get("separator", "    ")
    parts = []
    for pair in block.get("pairs", []):
        label = pair.get("label", "")
        value = pair.get("value", "")
        parts.append(f"{label}: {value}" if label else str(value))
    paragraph = doc.add_paragraph(separator.join(parts))
    apply_paragraph_style(paragraph, block)


def add_table_block(doc: Any, block: dict[str, Any]) -> None:
    from docx.shared import Inches

    headers = block.get("headers")
    rows = block.get("rows", [])
    col_count = len(headers) if headers else (len(rows[0]) if rows else 1)
    table_rows = len(rows) + (1 if headers else 0)
    table = doc.add_table(rows=table_rows, cols=col_count)
    table.style = "Table Grid"

    start_row = 0
    if headers:
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = str(header)
        start_row = 1

    for row_idx, row in enumerate(rows, start=start_row):
        for col_idx, value in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = str(value)

    widths = block.get("column_widths")
    if widths and len(widths) == col_count:
        total = float(sum(widths)) or 1.0
        for i, width in enumerate(widths):
            table.columns[i].width = Inches(6.5 * (width / total))

    align = block.get("align")
    if align:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = get_docx_alignment({"align": align})

    doc.add_paragraph("")


def render_docx_block(doc: Any, block: dict[str, Any]) -> None:
    block_type = block.get("type")
    if block_type == "paragraph":
        add_paragraph_block(doc, block)
    elif block_type in {"field_row", "labeled_row"}:
        add_field_row_block(doc, block)
    elif block_type == "table":
        add_table_block(doc, block)
    elif block_type == "spacer":
        add_spacer(doc, block.get("lines", 1))
    elif block_type == "page_break":
        doc.add_page_break()
    elif block_type in {"seal", "footer", "note"}:
        add_paragraph_block(doc, block)
    elif block_type == "heading":
        level = min(max(block.get("level", 2), 1), 3)
        doc.add_heading(block.get("text", ""), level=level)
    elif block_type == "bullets":
        for item in block.get("items", []):
            doc.add_paragraph(str(item), style="List Bullet")
    else:
        text = block.get("text")
        if text:
            add_paragraph_block(doc, block)


def configure_page(doc: Any, page: dict[str, Any]) -> None:
    section = doc.sections[0]
    orientation = page.get("orientation", "portrait").lower()
    if orientation == "landscape":
        from docx.enum.section import WD_ORIENT

        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    margin_map = {
        "margin_top_mm": "top_margin",
        "margin_bottom_mm": "bottom_margin",
        "margin_left_mm": "left_margin",
        "margin_right_mm": "right_margin",
    }
    from docx.shared import Mm

    for key, attr in margin_map.items():
        if key in page:
            setattr(section, attr, Mm(page[key]))


def add_certification(doc: Any, statement: dict[str, Any]) -> None:
    doc.add_page_break()
    heading = doc.add_heading("CERTIFICATE OF ACCURATE TRANSLATION", level=1)
    heading.alignment = get_docx_alignment({"align": "center"})
    doc.add_paragraph(statement.get("body", ""))
    for label, key in (
        ("Translator", "translator"),
        ("Date", "date"),
        ("Contact", "contact"),
    ):
        value = statement.get(key)
        if value:
            doc.add_paragraph(f"{label}: {value}")


def build_docx(data: dict[str, Any], output_path: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit("Missing dependency: pip install python-docx") from exc

    doc = Document()
    page = data.get("page")
    if page:
        configure_page(doc, page)

    mirror = data.get("mirror_source_layout", True)
    suppress_title = data.get("suppress_default_title", mirror)

    if not suppress_title:
        title = data.get("title", "OFFICIAL TRANSLATION")
        doc.add_heading(title, 0)
        doc_type = data.get("document_type")
        if doc_type:
            paragraph = doc.add_paragraph(doc_type)
            paragraph.alignment = get_docx_alignment({"align": "center"})

    for block in data.get("sections", []):
        render_docx_block(doc, block)

    statement = data.get("certification_statement")
    if statement:
        add_certification(doc, statement)

    doc.save(output_path)


def write_pdf_line(
    pdf: Any,
    text: str,
    *,
    height: float = 6,
    bold: bool = False,
    size: int = 11,
    align: str = "left",
) -> None:
    pdf.set_font("Helvetica", "B" if bold else "", size)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(pdf.epw, height, text, align=align)


def render_pdf_block(pdf: Any, block: dict[str, Any]) -> None:
    block_type = block.get("type")
    align = block.get("align", "left")
    size = block.get("font_size", 11)
    bold = bool(block.get("bold"))

    if block_type == "paragraph":
        write_pdf_line(pdf, block.get("text", ""), height=6, bold=bold, size=size, align=align)
        pdf.ln(block.get("space_after", 2) / 2 or 2)
    elif block_type in {"field_row", "labeled_row"}:
        separator = block.get("separator", "    ")
        parts = []
        for pair in block.get("pairs", []):
            label = pair.get("label", "")
            value = pair.get("value", "")
            parts.append(f"{label}: {value}" if label else str(value))
        write_pdf_line(pdf, separator.join(parts), bold=bold, size=size, align=align)
        pdf.ln(2)
    elif block_type == "table":
        headers = block.get("headers")
        rows = block.get("rows", [])
        if headers:
            write_pdf_line(pdf, " | ".join(str(h) for h in headers), bold=True, size=max(size - 1, 9), align=align)
        for row in rows:
            write_pdf_line(pdf, " | ".join(str(v) for v in row), size=max(size - 1, 9), align=align)
        pdf.ln(3)
    elif block_type == "spacer":
        pdf.ln(max(1, block.get("lines", 1)) * 4)
    elif block_type == "page_break":
        pdf.add_page()
    elif block_type in {"seal", "footer", "note"}:
        if block.get("space_before"):
            pdf.ln(block["space_before"] / 2)
        write_pdf_line(
            pdf,
            block.get("text", ""),
            bold=bold,
            size=block.get("font_size", size),
            align=align,
        )
        pdf.ln(2)
    elif block_type == "heading":
        write_pdf_line(pdf, block.get("text", ""), bold=True, size=size + 2, align=align)
        pdf.ln(2)
    elif block_type == "bullets":
        for item in block.get("items", []):
            write_pdf_line(pdf, f"- {item}", size=size, align=align)
        pdf.ln(2)


def build_pdf_with_fpdf(data: dict[str, Any], output_path: Path) -> None:
    try:
        from fpdf import FPDF
    except ImportError as exc:
        raise SystemExit("Missing dependency: pip install fpdf2") from exc

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    page = data.get("page", {})
    if page.get("orientation", "portrait").lower() == "landscape":
        pdf = FPDF(orientation="L")
        pdf.set_auto_page_break(auto=True, margin=15)

    pdf.add_page()

    mirror = data.get("mirror_source_layout", True)
    suppress_title = data.get("suppress_default_title", mirror)

    if not suppress_title:
        write_pdf_line(
            pdf,
            data.get("title", "OFFICIAL TRANSLATION"),
            height=10,
            bold=True,
            size=16,
            align="center",
        )
        pdf.ln(4)
        doc_type = data.get("document_type")
        if doc_type:
            write_pdf_line(pdf, doc_type, height=8, size=12, align="center")
            pdf.ln(4)

    for block in data.get("sections", []):
        render_pdf_block(pdf, block)

    statement = data.get("certification_statement")
    if statement:
        pdf.add_page()
        write_pdf_line(
            pdf,
            "CERTIFICATE OF ACCURATE TRANSLATION",
            height=8,
            bold=True,
            size=14,
            align="center",
        )
        pdf.ln(4)
        write_pdf_line(pdf, statement.get("body", ""))
        pdf.ln(4)
        for label, key in (
            ("Translator", "translator"),
            ("Date", "date"),
            ("Contact", "contact"),
        ):
            value = statement.get(key)
            if value:
                write_pdf_line(pdf, f"{label}: {value}")

    pdf.output(str(output_path))


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    if shutil.which("soffice"):
        out_dir = pdf_path.parent
        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(docx_path),
            ],
            capture_output=True,
            text=True,
        )
        generated = out_dir / f"{docx_path.stem}.pdf"
        if result.returncode == 0 and generated.exists():
            if generated != pdf_path:
                generated.replace(pdf_path)
            return True

    if shutil.which("pandoc"):
        result = subprocess.run(
            ["pandoc", str(docx_path), "-o", str(pdf_path)],
            capture_output=True,
            text=True,
        )
        if result.returncode == 0 and pdf_path.exists():
            return True

    return False


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit(f"Usage: {sys.argv[0]} translation.json")

    input_path = Path(sys.argv[1]).resolve()
    data = load_input(input_path)

    output_dir = Path(data.get("output_dir", input_path.parent)).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    basename = data.get("basename", input_path.stem.replace("_translation", ""))
    docx_path = output_dir / f"{basename}.docx"
    pdf_path = output_dir / f"{basename}.pdf"

    build_docx(data, docx_path)

    if not convert_docx_to_pdf(docx_path, pdf_path):
        build_pdf_with_fpdf(data, pdf_path)

    print(f"DOCX: {docx_path}")
    print(f"PDF:  {pdf_path}")


if __name__ == "__main__":
    main()
